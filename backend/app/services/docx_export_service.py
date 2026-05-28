from io import BytesIO

from docx import Document

from app.db.models import Question, Quiz


class QuizDocxExportService:
    def build_quiz_docx(self, quiz: Quiz, questions: list[Question]) -> bytes:
        document = Document()
        document.add_heading(quiz.title or "Quiz", level=0)

        meta = (
            f"Subject: {quiz.subject or '-'} | "
            f"Grade: {quiz.grade or '-'} | "
            f"Difficulty: {quiz.difficulty or '-'}"
        )
        document.add_paragraph(meta)

        for index, question in enumerate(questions, start=1):
            document.add_heading(f"Question {index}", level=2)
            document.add_paragraph(f"[{question.question_type}] {question.question_text}")

            for option in question.answers or []:
                marker = "✓" if option in (question.correct_answers or []) else "-"
                document.add_paragraph(f"{marker} {option}", style="List Bullet")

            if question.explanation:
                document.add_paragraph(f"Explanation: {question.explanation}")

        buffer = BytesIO()
        document.save(buffer)
        return buffer.getvalue()


docx_export_service = QuizDocxExportService()
